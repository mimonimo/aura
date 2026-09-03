"""계획서 초안 내보내기 — 마크다운 초안을 실제 문서 파일로.

초안 계약: "## 섹션명" 헤딩 + 문단 + 파이프 표(마크다운). 이 모듈이 그 계약을
파싱해 docx(한글에서도 열림)·pdf·md로 만든다. HWPX 직접 생성은 라이브러리
검토 후 추가 예정 — 그 전까지 docx가 한글 호환 경로다.
"""

from __future__ import annotations

import io
import re

_HEAD = re.compile(r"^##\s+(.+)$")
_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")


def parse_draft(draft: str) -> list[dict]:
    """초안 → [{title, blocks:[{kind:'p'|'table', text|rows}]}] 구조."""
    sections: list[dict] = []
    cur: dict | None = None
    lines = draft.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        m = _HEAD.match(line)
        if m:
            cur = {"title": m.group(1).strip(), "blocks": []}
            sections.append(cur)
            i += 1
            continue
        if cur is None:
            cur = {"title": "", "blocks": []}
            sections.append(cur)
        if line.startswith("|") and line.endswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                i += 1
                if _TABLE_SEP.match(row_line):
                    continue  # |---|---| 구분줄
                rows.append([c.strip() for c in row_line.strip("|").split("|")])
            if rows:
                cur["blocks"].append({"kind": "table", "rows": rows})
            continue
        # 문단 — 빈 줄 전까지 이어붙인다
        if line.strip():
            para = [line.strip()]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(("|", "## ")):
                para.append(lines[i].strip())
                i += 1
            cur["blocks"].append({"kind": "p", "text": "\n".join(para)})
            continue
        i += 1
    return [s for s in sections if s["blocks"] or s["title"]]


def build_draft_docx(
    title: str, draft: str, images: list[str] | None = None
) -> bytes:
    """초안 → .docx (제목·섹션 헤딩·문단·실제 표·붙임 그림)."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(title, level=0)
    for sec in parse_draft(draft):
        if sec["title"]:
            doc.add_heading(sec["title"], level=1)
        for b in sec["blocks"]:
            if b["kind"] == "table":
                rows = b["rows"]
                n_cols = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=n_cols)
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.rows[ri].cells[ci].text = cell
                doc.add_paragraph("")
            else:
                doc.add_paragraph(b["text"])
    if images:
        doc.add_heading("붙임 그림", level=1)
        for img in images:
            try:
                doc.add_picture(img, width=Inches(5.5))
            except Exception:
                continue
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_draft_hwpx(
    title: str, draft: str, images: list[str] | None = None
) -> bytes | None:
    """초안 → .hwpx (python-hwpx — 순수 파이썬 OWPML 생성, 한글에서 열림)."""
    try:
        from hwpx import HwpxDocument

        doc = HwpxDocument.new()
        # 한글의 '개요' 스타일은 기본 개요 번호(가, 나, 다…)를 자동으로 붙여
        # 제목이 "가. 요약"처럼 보인다 — 개요 스타일 대신 바탕글 문단에
        # 우리가 직접 번호를 매긴다
        doc.add_paragraph(title, style="바탕글")
        doc.add_paragraph("", style="바탕글")
        for i, sec in enumerate(parse_draft(draft), start=1):
            if sec["title"]:
                doc.add_paragraph("", style="바탕글")
                doc.add_paragraph(f"{i}. {sec['title']}", style="바탕글")
            for b in sec["blocks"]:
                if b["kind"] == "table":
                    rows = b["rows"]
                    n_cols = max(len(r) for r in rows)
                    t = doc.add_table(rows=len(rows), cols=n_cols)
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            t.set_cell_text(ri, ci, cell, logical=True)
                else:
                    for para in b["text"].split("\n"):
                        doc.add_paragraph(para, style="바탕글")
        if images:
            from pathlib import Path as _P

            doc.add_paragraph("", style="바탕글")
            doc.add_paragraph("붙임 그림", style="바탕글")
            for img in images:
                try:
                    fmt = _P(img).suffix.lstrip(".").lower() or "png"
                    doc.add_picture(
                        _P(img).read_bytes(), fmt, width_mm=140.0
                    )
                except Exception:
                    continue
        buf = io.BytesIO()
        doc.save_to_stream(buf)
        return buf.getvalue()
    except Exception:
        import logging

        logging.getLogger(__name__).warning("hwpx 내보내기 실패", exc_info=True)
        return None


def build_draft_pdf(
    title: str, draft: str, images: list[str] | None = None
) -> bytes | None:
    """초안 → .pdf (Platypus 흐름 배치 — 제목·문단·표)."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        font = "HYSMyeongJo-Medium"
        try:
            pdfmetrics.getFont(font)
        except KeyError:
            pdfmetrics.registerFont(UnicodeCIDFont(font))

        st_title = ParagraphStyle("t", fontName=font, fontSize=17, leading=22)
        st_head = ParagraphStyle("h", fontName=font, fontSize=13, leading=18,
                                 spaceBefore=12, spaceAfter=4)
        st_body = ParagraphStyle("b", fontName=font, fontSize=10.5, leading=16)

        buf = io.BytesIO()
        pdf = SimpleDocTemplate(buf, pagesize=A4, title=title,
                                topMargin=52, bottomMargin=48)
        story = [Paragraph(title, st_title), Spacer(1, 14)]
        for sec in parse_draft(draft):
            if sec["title"]:
                story.append(Paragraph(sec["title"], st_head))
            for b in sec["blocks"]:
                if b["kind"] == "table":
                    rows = [
                        [Paragraph(c, st_body) for c in r] for r in b["rows"]
                    ]
                    t = Table(rows, hAlign="LEFT")
                    t.setStyle(TableStyle([
                        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#8A94A6")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2F8")),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    story += [Spacer(1, 4), t, Spacer(1, 8)]
                else:
                    story.append(Paragraph(
                        b["text"].replace("\n", "<br/>"), st_body))
                    story.append(Spacer(1, 6))
        if images:
            from reportlab.platypus import Image as RLImage

            story.append(Paragraph("붙임 그림", st_head))
            for img in images:
                try:
                    from PIL import Image as PILImage

                    with PILImage.open(img) as im:
                        w, h = im.size
                    max_w = 430.0
                    scale_f = min(max_w / w, 1.0)
                    story += [
                        RLImage(img, width=w * scale_f, height=h * scale_f),
                        Spacer(1, 8),
                    ]
                except Exception:
                    continue
        pdf.build(story)
        return buf.getvalue()
    except Exception:
        return None
