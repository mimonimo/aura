"""추출 결과 렌더링 — doc_chunks를 문서 모양의 HTML 블록으로 바꾼다.

표 조각은 JSON 구조({n_rows, n_cols, cells})로 저장되며, 병합 셀(rowspan/colspan)과
머리글(th)을 유지한 채 실제 표로 그린다. 셀 내용은 전부 이스케이프한다.
"""

from __future__ import annotations

import json
import re
from collections import defaultdict

from markupsafe import Markup, escape


def _rich(text: str) -> Markup:
    """이스케이프 후 **굵게**만 살린다 — 비전 판독의 강조 보존."""
    escaped = str(escape(text))
    return Markup(re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped))


_TOC_LINE = re.compile(r"[·.…]{3,}\s*\d{1,3}")


def _toc_html(rows: dict[int, list], n_rows: int) -> Markup:
    """목차 표 → 항목·쪽번호 목록. 2단 목차는 열 우선(왼쪽 열 먼저)으로 읽는다."""
    cols: dict[int, list[str]] = defaultdict(list)
    for r in range(n_rows):
        for c, _rs, _cs, _hd, txt in sorted(rows.get(r, [])):
            if str(txt).strip():
                cols[c].append(str(txt))
    items: list[tuple[str, str]] = []
    for c in sorted(cols):
        for cell in cols[c]:
            # 셀 안에 여러 항목이 붙어 있으면 '제목 ···· 쪽수' 단위로 자른다
            for m in re.finditer(r"(.+?)[·.…]{3,}\s*(\d{1,3})", cell):
                title = m.group(1).strip(" ·.…-")
                if title:
                    items.append((title, m.group(2)))
    if not items:
        return Markup('<pre class="doc-text">{}</pre>').format(
            "\n".join(t for col in cols.values() for t in col)
        )
    parts = ['<div class="extract-toc"><div class="extract-toc-title">목차</div>']
    for title, page in items:
        parts.append(
            f'<div class="toc-row"><span class="toc-t">{escape(title)}</span>'
            f'<span class="toc-p">{escape(page)}</span></div>'
        )
    parts.append("</div>")
    return Markup("".join(parts))


def table_html(content: str) -> Markup:
    try:
        data = json.loads(content)
        cells = data["cells"]
        n_rows = int(data["n_rows"])
    except (json.JSONDecodeError, KeyError, TypeError, ValueError):
        # 구조를 못 읽으면 원문 그대로 (구버전 조각 호환)
        return Markup('<pre class="doc-text">{}</pre>').format(content)

    rows: dict[int, list] = defaultdict(list)
    for r, c, rs, cs, hd, txt in cells:
        rows[int(r)].append((int(c), int(rs), int(cs), bool(hd), str(txt)))

    # 목차 표 감지 — 점선 리더(····)+쪽번호 셀이 많으면 표 대신 목차로 그린다
    all_texts = [str(t) for *_, t in cells if str(t).strip()]
    n_toc = sum(1 for t in all_texts if _TOC_LINE.search(t))
    if all_texts and n_toc * 2 >= len(all_texts):
        return _toc_html(rows, int(data["n_rows"]))

    parts = ['<div class="table-scroll"><table class="extract">']
    for r in range(n_rows):
        parts.append("<tr>")
        for c, rs, cs, hd, txt in sorted(rows.get(r, [])):
            tag = "th" if hd else "td"
            attrs = (f' rowspan="{rs}"' if rs > 1 else "") + (
                f' colspan="{cs}"' if cs > 1 else ""
            )
            parts.append(f"<{tag}{attrs}>{escape(txt)}</{tag}>")
        parts.append("</tr>")
    parts.append("</table></div>")
    return Markup("".join(parts))


def table_csv(content: str) -> str:
    """표 JSON → CSV (엑셀 호환, 병합 셀은 좌상단 셀에만 값)."""
    import csv
    import io

    data = json.loads(content)
    grid = [["" for _ in range(int(data["n_cols"]))] for _ in range(int(data["n_rows"]))]
    for r, c, _rs, _cs, _hd, txt in data["cells"]:
        if int(r) < len(grid) and int(c) < len(grid[0]):
            grid[int(r)][int(c)] = str(txt)
    buf = io.StringIO()
    csv.writer(buf).writerows(grid)
    return buf.getvalue()


def export_markdown(filename: str, chunks: list[dict]) -> str:
    """추출 결과 전체를 마크다운으로 — 소제목·문단·표(파이프 그리드)."""
    lines = [f"# {filename} — 추출 결과", ""]
    for c in chunks:
        if c["kind"] == "heading":
            lines += [f"### {c['content']}", ""]
        elif c["kind"] == "table":
            try:
                data = json.loads(c["content"])
            except (json.JSONDecodeError, TypeError):
                lines += [c["content"], ""]
                continue
            grid = [
                ["" for _ in range(int(data["n_cols"]))]
                for _ in range(int(data["n_rows"]))
            ]
            for r, col, _rs, _cs, _hd, txt in data["cells"]:
                if int(r) < len(grid) and int(col) < len(grid[0]):
                    grid[int(r)][int(col)] = str(txt).replace("\n", " ")
            if grid:
                lines.append("| " + " | ".join(grid[0]) + " |")
                lines.append("|" + "---|" * len(grid[0]))
                for row in grid[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
        else:
            lines += [c["content"], ""]
    return "\n".join(lines)


def chunk_blocks(
    chunks: list[dict],
    doc_id: int | None = None,
    asset_by_name: dict[str, int] | None = None,
) -> list[Markup]:
    """조각 목록을 순서대로 HTML 블록으로 — 문단·표·그림이 원래 순서·페이지대로 흐른다."""
    blocks: list[Markup] = []
    last_page: int | None = None
    for c in chunks:
        pg = c.get("page_no")
        if pg and pg != last_page:
            if last_page is not None:
                blocks.append(
                    Markup('<div class="extract-page">{}쪽</div>').format(pg)
                )
            last_page = pg
        if c["kind"] == "image":
            aid = (asset_by_name or {}).get(c["content"])
            if doc_id is not None and aid:
                blocks.append(Markup(
                    '<figure style="margin:6px 0 16px;">'
                    '<img src="/doc/{d}/asset/{a}" style="max-width:70%; border:1px solid'
                    ' var(--line); border-radius:10px; display:block;">'
                    '<figcaption class="muted" style="font-size:11px; margin-top:3px;">'
                    '추출 그림 · <a href="/doc/{d}/asset/{a}?dl=1">내려받기</a>'
                    '</figcaption></figure>'
                ).format(d=doc_id, a=aid))
            continue
        if c["kind"] == "table":
            block = table_html(c["content"])
            if doc_id is not None and c.get("id"):
                block += Markup(
                    '<p style="margin:-10px 0 14px; text-align:right;">'
                    '<a href="/doc/{}/table/{}.csv" class="muted"'
                    ' style="font-size:11.5px;">표 CSV 내려받기</a></p>'
                ).format(doc_id, c["id"])
            blocks.append(block)
        elif c["kind"] == "heading":
            blocks.append(Markup('<h4 class="extract-h">{}</h4>').format(_rich(c["content"])))
        else:
            blocks.append(Markup('<p class="extract-p">{}</p>').format(_rich(c["content"])))
    return blocks


def layout_pages(
    chunks: list[dict],
    doc_id: int | None = None,
    asset_by_name: dict[str, int] | None = None,
    page_sizes: dict[int, tuple[float, float]] | None = None,
) -> list[Markup] | None:
    """bbox로 원본 배치를 재구성 — 실제 페이지 치수 기준, 픽셀 고정 캔버스.

    글자 크기는 박스 높이·줄 수에서 계산해 원본 밀도에 가깝게 맞춘다.
    좌표가 있는 조각이 충분치 않으면 None.
    """
    CANVAS_W = 760.0
    with_bbox = [c for c in chunks if c.get("bbox")]
    if len(with_bbox) < max(2, len(chunks) // 3):
        return None

    pages: dict[int, list] = defaultdict(list)
    for c in chunks:
        if not c.get("bbox"):
            continue
        try:
            x0, y0, x1, y1 = (float(v) for v in c["bbox"].split(","))
        except ValueError:
            continue
        if x1 <= x0 or y1 <= y0:
            continue
        pages[c.get("page_no") or 1].append((c, x0, y0, x1, y1))

    out: list[Markup] = []
    for pg in sorted(pages):
        boxes = pages[pg]
        if page_sizes and pg in page_sizes:
            pw, ph = page_sizes[pg]
        else:
            pw = max(x1 for *_, x1, _ in boxes) * 1.02
            ph = max(y1 for *_, y1 in boxes) * 1.03
        if pw <= 0 or ph <= 0:
            continue
        # 좌표계 보정 — bbox가 렌더 픽셀 기준(페이지의 2배 등)이면 균일 축소
        max_x = max(x1 for *_, x1, _ in boxes)
        max_y = max(y1 for *_, y1 in boxes)
        fit = max(max_x / pw, max_y / ph, 1.0)
        if fit > 1.08:
            boxes = [
                (c, x0 / fit, y0 / fit, x1 / fit, y1 / fit)
                for c, x0, y0, x1, y1 in boxes
            ]
        scale = CANVAS_W / pw
        parts = [
            f'<div class="layout-page" style="width:{CANVAS_W:.0f}px;'
            f' height:{ph * scale:.0f}px;">',
            f'<span class="layout-pageno">{pg}쪽</span>',
        ]
        for c, x0, y0, x1, y1 in boxes:
            left, top = x0 * scale, y0 * scale
            bw, bh = (x1 - x0) * scale, (y1 - y0) * scale
            style = (
                f"left:{left:.1f}px; top:{top:.1f}px;"
                f" width:{bw:.1f}px; height:{bh:.1f}px;"
            )
            if c["kind"] == "image":
                aid = (asset_by_name or {}).get(c["content"])
                inner = (
                    Markup('<img src="/doc/{}/asset/{}" style="width:100%; height:100%;'
                           ' object-fit:contain;">').format(doc_id, aid)
                    if doc_id is not None and aid else Markup("")
                )
            elif c["kind"] == "table":
                inner = table_html(c["content"])
            else:
                n_lines = max(c["content"].count("\n") + 1, 1)
                fs = min(15.0, bh / n_lines * 0.72)
                # 폭 기준 상한 — 가장 긴 줄이 박스 폭을 넘지 않게 (CJK≈1em, 그 외≈0.55em)
                widest = max(
                    (sum(1.0 if ord(ch) > 0x2E80 else 0.55 for ch in ln)
                     for ln in c["content"].splitlines() if ln.strip()),
                    default=1.0,
                )
                fs = max(4.5, min(fs, bw / max(widest, 1.0) * 0.96))
                weight = "700" if c["kind"] == "heading" else "400"
                color = "var(--navy)" if c["kind"] == "heading" else "inherit"
                inner = Markup(
                    '<span style="font-size:{:.1f}px; font-weight:{}; color:{};'
                    ' line-height:1.32; display:block;">{}</span>'
                ).format(fs, weight, color, _rich(c["content"]))
            parts.append(
                Markup('<div class="layout-box" style="{}">{}</div>').format(
                    Markup(style), inner
                )
            )
        parts.append("</div>")
        out.append(Markup("".join(str(x) for x in parts)))
    return out or None


def trailing_image_blocks(doc_id: int, assets: list[dict]) -> list[Markup]:
    """위치 정보가 없는 추출 그림을 프리뷰 말미 섹션으로 — 복원 문서와 동일 구성."""
    if not assets:
        return []
    blocks = [Markup('<h4 class="extract-h">추출 그림</h4>')]
    for a in assets:
        blocks.append(Markup(
            '<figure style="margin:6px 0 16px;">'
            '<img src="/doc/{d}/asset/{a}" style="max-width:70%; border:1px solid'
            ' var(--line); border-radius:10px; display:block;">'
            '<figcaption class="muted" style="font-size:11px; margin-top:3px;">'
            'p{p} · <a href="/doc/{d}/asset/{a}?dl=1">내려받기</a>'
            '</figcaption></figure>'
        ).format(d=doc_id, a=a["id"], p=a.get("page_no") or "?"))
    return blocks


def build_docx(
    filename: str,
    chunks: list[dict],
    asset_paths: dict[str, str] | None = None,
    extra_images: list[str] | None = None,
) -> bytes:
    """추출 조각을 편집 가능한 워드 문서로 복원 — 제목·문단·병합 표·그림.

    전자 문서 복원이 이 파이프라인의 최종 목표다 (ADR-0007).
    """
    import io

    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    for c in chunks:
        kind = c["kind"]
        if kind == "heading":
            doc.add_heading(c["content"], level=2)
        elif kind == "table":
            try:
                data = json.loads(c["content"])
                n_rows, n_cols = int(data["n_rows"]), int(data["n_cols"])
                if n_rows < 1 or n_cols < 1:
                    continue
                t = doc.add_table(rows=n_rows, cols=n_cols)
                t.style = "Table Grid"
                for r, col, rs, cs, hd, txt in data["cells"]:
                    r, col, rs, cs = int(r), int(col), int(rs), int(cs)
                    if r >= n_rows or col >= n_cols:
                        continue
                    cell = t.cell(r, col)
                    end_r = min(r + rs - 1, n_rows - 1)
                    end_c = min(col + cs - 1, n_cols - 1)
                    if (end_r, end_c) != (r, col):
                        cell = cell.merge(t.cell(end_r, end_c))
                    cell.text = str(txt)
                    if hd:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                doc.add_paragraph("")
            except (json.JSONDecodeError, KeyError, TypeError, ValueError):
                doc.add_paragraph(c["content"])
        elif kind == "image":
            path = (asset_paths or {}).get(c["content"])
            if path:
                try:
                    doc.add_picture(path, width=Inches(5.2))
                except Exception:
                    pass
        else:
            # **굵게** 강조를 워드 굵기로 옮긴다
            para = doc.add_paragraph()
            for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", c["content"])):
                run = para.add_run(part)
                run.bold = i % 2 == 1

    # 본문에 그림 위치가 없는 경우(비전 전사 등) — 추출 그림을 끝에 첨부한다
    placed = any(c["kind"] == "image" for c in chunks)
    if not placed and extra_images:
        doc.add_heading("추출 그림", level=2)
        for path in extra_images[:20]:
            try:
                doc.add_picture(path, width=Inches(5.2))
            except Exception:
                continue

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
