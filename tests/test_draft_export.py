"""초안 내보내기 — 사용자 요구 항목 구조의 초안을 docx/pdf/md 파일로."""
import io

DRAFT = """## 사업 개요
본 사업은 지역 인재 양성을 목표로 한다.

## 예산 계획
| 항목 | 단가 | 수량 | 금액 |
|---|---|---|---|
| 장학금 | 250만 원 | 24명 | 6,000만 원 |
| 운영비 | 50만 원 | 10회 | 500만 원 |

## 추진 일정
1분기에 착수한다."""


def test_parse_draft_sections_and_tables():
    from zzaimy.app.draft_export import parse_draft

    secs = parse_draft(DRAFT)
    assert [s["title"] for s in secs] == ["사업 개요", "예산 계획", "추진 일정"]
    table_blocks = [b for b in secs[1]["blocks"] if b["kind"] == "table"]
    assert len(table_blocks) == 1
    assert table_blocks[0]["rows"][0] == ["항목", "단가", "수량", "금액"]
    assert table_blocks[0]["rows"][1][3] == "6,000만 원"


def test_draft_docx_contains_headings_and_table():
    from docx import Document

    from zzaimy.app.draft_export import build_draft_docx

    payload = build_draft_docx("테스트 계획서", DRAFT)
    doc = Document(io.BytesIO(payload))
    heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "사업 개요" in heads and "예산 계획" in heads
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[1].cells[0].text == "장학금"


def test_draft_pdf_text_extractable():
    from pypdf import PdfReader

    from zzaimy.app.draft_export import build_draft_pdf

    payload = build_draft_pdf("테스트 계획서", DRAFT)
    assert payload is not None
    text = "".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(payload)).pages)
    assert "사업 개요" in text and "지역 인재 양성" in text
    assert "6,000만 원" in text  # 표 내용도 문서에 들어간다
